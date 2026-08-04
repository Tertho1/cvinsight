"""
scripts/generate_benchmark_cvs.py

Generate a controlled benchmark set of demo CVs (demo/benchmark/) covering the
failure scenarios found in the extraction audit. Reproducible: run this script
and the same files are produced every time.

Each scenario maps to a documented extraction issue (docs/extraction_audit.md,
docs/extraction_improvements.md). Content is realistic, uses real dates, and
skills are drawn from config/skill_taxonomy.json so extraction has a fair
chance.

Usage:
    python scripts/generate_benchmark_cvs.py
"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, "demo", "benchmark")
TAXONOMY = json.load(open(os.path.join(ROOT, "config", "skill_taxonomy.json"), encoding="utf-8"))
CATS = TAXONOMY["categories"]


def pick(*cats):
    out = []
    for c in cats:
        out.extend(CATS.get(c, []))
    return out


def clean_json(x):
    if isinstance(x, dict):
        return {k: clean_json(v) for k, v in x.items()}
    if isinstance(x, list):
        return [clean_json(v) for v in x]
    if isinstance(x, str):
        return x.replace("\ufffd", "-")
    return x


def save_docx(doc, filename):
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    print(f"  wrote {filename}")


def add_docx_table(doc, rows, cols=None, duplicate=False):
    """Add a table where each cell holds one line of content.

    duplicate=True emits each row twice (merged-cell DOCX pattern).
    """
    ncols = cols or max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.autofit = True
    for row in rows:
        cells = row + [""] * (ncols - len(row))
        tr = table.add_row()
        for i, val in enumerate(cells[:ncols]):
            tr.cells[i].text = val
        if duplicate:
            tr2 = table.add_row()
            for i, val in enumerate(cells[:ncols]):
                tr2.cells[i].text = val
    return table


# ---------------------------------------------------------------------------
# Scenario 1: clean standard DOCX (control / baseline)
# ---------------------------------------------------------------------------
def scenario_clean():
    doc = Document()
    doc.add_heading("Ava Robinson", level=0)
    doc.add_paragraph("ava.robinson@example.com | (617) 555-0142 | Boston, MA")
    doc.add_paragraph("Senior Software Engineer with 8 years of experience building scalable web applications.")

    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Senior Software Engineer, CloudScale Inc | Jan 2021 - Present")
    doc.add_paragraph("- Lead a team of 5 engineers building microservices in Python and Go")
    doc.add_paragraph("- Designed event-driven architecture using Kafka and AWS Lambda")
    doc.add_paragraph("Software Engineer, DataWorks LLC | Mar 2018 - Dec 2020")
    doc.add_paragraph("- Built REST APIs with Django and PostgreSQL")
    doc.add_paragraph("- Migrated legacy monolith to Kubernetes")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("M.Sc. Computer Science, MIT, 2018")
    doc.add_paragraph("B.Sc. Software Engineering, University of Washington, 2016")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Go, JavaScript, Django, React, PostgreSQL, MongoDB, Docker, Kubernetes, AWS, Kafka, REST, Git, CI/CD")

    doc.add_heading("Projects", level=1)
    doc.add_paragraph("Real-time Analytics Dashboard - React, FastAPI, Redis | https://github.com/ava/analytics")
    doc.add_paragraph("Event Sourcing Library - Python, Kafka")

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Certified Solutions Architect - Associate, 2022")

    doc.add_heading("Languages", level=1)
    doc.add_paragraph("English (Native), Spanish (Conversational)")

    doc.add_heading("Leadership", level=1)
    doc.add_paragraph("Mentor, Women in Tech Boston")
    save_docx(doc, "01_clean_standard.docx")


# ---------------------------------------------------------------------------
# Scenario 2: two-column PDF layout (sidebar + main column)
# ---------------------------------------------------------------------------
class TwoColumnPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def sidebar(self):
        self.set_font("Helvetica", "B", 12)
        self.set_xy(10, 10)
        self.cell(60, 8, "PROFILE", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 9)
        self.set_xy(10, 20)
        self.multi_cell(55, 5, "Backend engineer focused on data pipelines and cloud infrastructure. 6 years of experience.")

    def header_block(self, text):
        self.set_font("Helvetica", "B", 15)
        self.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")


def scenario_two_column_pdf():
    pdf = TwoColumnPDF()
    pdf.add_page()
    pdf.set_xy(10, 10)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Owen Bennett", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 6, "owen.bennett@example.com | (312) 555-0198 | Chicago, IL",
             new_x="LMARGIN", new_y="NEXT")

    # Main column (right of a ~62pt gutter)
    main_x, main_w = 78, 122
    pdf.set_xy(main_x, 30)
    pdf.set_font("Helvetica", "B", 12)
    pdf.multi_cell(main_w, 6, "EXPERIENCE")

    jobs = [
        ("Senior Data Engineer, Lakehouse Analytics Inc | Apr 2020 - Present",
         "Designed batch and streaming pipelines using Spark, Airflow, and Kafka.\nReduced pipeline cost 30% by optimizing storage tiering in S3."),
        ("Data Engineer, FinStreet Corp | Jun 2017 - Mar 2020",
         "Built ETL jobs in Python and SQL Server.\nIntroduced dbt-based transformation layer."),
        ("Software Engineer, Nebula Solutions | Jul 2015 - May 2017",
         "Developed internal tools with Python and PostgreSQL."),
    ]
    y = 36
    for title, body in jobs:
        pdf.set_xy(main_x, y)
        pdf.set_font("Helvetica", "B", 10)
        pdf.multi_cell(main_w, 5, title)
        y = pdf.get_y() + 1
        pdf.set_xy(main_x, y)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(main_w, 5, body)
        y = pdf.get_y() + 3

    pdf.set_xy(main_x, y + 3)
    pdf.set_font("Helvetica", "B", 12)
    pdf.multi_cell(main_w, 6, "EDUCATION")
    pdf.set_xy(main_x, pdf.get_y() + 1)
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(main_w, 5, "B.Sc. Computer Science, University of Illinois at Chicago, 2015")

    # Sidebar (left column)
    pdf.set_xy(10, 78)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(55, 7, "CONTACT")
    pdf.set_font("Helvetica", "", 8)
    pdf.set_xy(10, 86)
    pdf.multi_cell(55, 4, "Chicago, IL\n+1 (312) 555-0198\nowen.bennett@example.com\nlinkedin.com/in/owenbennett")
    pdf.set_xy(10, pdf.get_y() + 3)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(55, 7, "SKILLS")
    pdf.set_font("Helvetica", "", 8)
    pdf.set_xy(10, pdf.get_y() + 1)
    pdf.multi_cell(55, 4, "Python, Spark, Airflow, Kafka, SQL, dbt, AWS, Docker, Git")

    path = os.path.join(OUT_DIR, "02_two_column.pdf")
    pdf.output(path)
    print("  wrote 02_two_column.pdf")


# ---------------------------------------------------------------------------
# Scenario 3: table-based DOCX (one entity per cell -> lines) with real dates
# ---------------------------------------------------------------------------
def scenario_table_docx():
    doc = Document()
    doc.add_heading("Maya Patel", level=0)
    doc.add_paragraph("maya.patel@example.com | (415) 555-0110 | San Francisco, CA")
    doc.add_paragraph("Full-stack developer focused on React and Node.js.")

    doc.add_heading("Experience", level=1)
    add_docx_table(doc, [
        ["Frontend Engineer", "Brightpath", "San Francisco, CA"],
        ["Jun 2021 - Present"],
        ["- Build React components for the marketing platform"],
        ["- Optimize page load time with code splitting"],
        ["Junior Developer", "Sparkbyte", "Oakland, CA"],
        ["Aug 2019 - May 2021"],
        ["- Maintained Node.js API and React dashboard"],
        ["- Wrote unit tests with Jest"],
    ])

    doc.add_heading("Education", level=1)
    add_docx_table(doc, [
        ["B.Sc. Computer Science", "University of California, Berkeley", "2019"],
    ])

    doc.add_heading("Skills", level=1)
    add_docx_table(doc, [
        ["React", "TypeScript", "Node.js", "Express"],
        ["PostgreSQL", "MongoDB", "Docker", "Git"],
    ])

    doc.add_heading("Projects", level=1)
    add_docx_table(doc, [
        ["Task Manager App", "React, Node.js, MongoDB"],
        ["Portfolio Site", "Next.js, Tailwind CSS"],
    ])
    save_docx(doc, "03_table_docx.docx")


# ---------------------------------------------------------------------------
# Scenario 4: duplicated merged-cell DOCX (the priya pattern)
# ---------------------------------------------------------------------------
def scenario_duplicated_docx():
    doc = Document()
    doc.add_heading("Daniel Chen", level=0)
    doc.add_paragraph("daniel.chen@example.com | (646) 555-0173 | New York, NY")
    doc.add_paragraph("Senior Web Developer specializing in front-end development.")

    doc.add_heading("Skill Highlights", level=1)
    for s in ["JavaScript", "React", "CSS", "SQL", "Project Management"]:
        doc.add_paragraph(s)

    doc.add_heading("Experience", level=1)
    add_docx_table(doc, [
        ["Web Developer - 09/2015 to 05/2019"],
        ["Brightpath Design, New York"],
        ["Cooperate with designers to create clean interfaces."],
        ["Develop project concepts and maintain optimal workflow."],
        ["Complete detailed programming and development tasks."],
    ], duplicate=True)

    doc.add_heading("Education", level=1)
    add_docx_table(doc, [
        ["Bachelor of Science: Computer Information Systems - 2014"],
        ["Columbia University, NY"],
    ], duplicate=True)

    doc.add_heading("Certifications", level=1)
    add_docx_table(doc, [
        ["PHP Framework (certificate): Zend, CodeIgniter, Symfony."],
        ["Programming Languages: JavaScript, HTML5, PHP, CSS, SQL."],
    ], duplicate=True)
    save_docx(doc, "04_duplicated_cells.docx")


# ---------------------------------------------------------------------------
# Scenario 5: academic CV (research/teaching experience, invited talks)
# ---------------------------------------------------------------------------
def scenario_academic():
    doc = Document()
    doc.add_heading("Prof. Elena Vasquez", level=0)
    doc.add_paragraph("elena.vasquez@university.edu | (202) 555-0166 | Washington, DC")
    doc.add_paragraph("Associate Professor of Computer Science.")

    doc.add_heading("Research Experience", level=1)
    doc.add_paragraph("Research Scientist, National AI Lab | Jan 2020 - Present")
    doc.add_paragraph("Led a team studying NLP robustness. Published 12 papers.")

    doc.add_heading("Teaching Experience", level=1)
    doc.add_paragraph("Associate Professor, Georgetown University | Sep 2016 - Present")
    doc.add_paragraph("Taught Machine Learning and Data Structures to 200+ students.")
    doc.add_paragraph("Teaching Assistant, Stanford University | Sep 2014 - May 2016")
    doc.add_paragraph("Ran discussion sections for Intro to Programming.")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("Ph.D. Computer Science, Stanford University, 2016")
    doc.add_paragraph("M.Sc. Computer Science, UC Berkeley, 2012")
    doc.add_paragraph("B.Sc. Mathematics, MIT, 2010")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("Vasquez, E., et al. Robust NLP under Domain Shift. ACL 2024.")

    doc.add_heading("Invited Talks", level=1)
    doc.add_paragraph("Keynote, ICML Workshop on Robustness, Jul 2024")
    doc.add_paragraph("Seminar, Carnegie Mellon University, Mar 2024")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, PyTorch, TensorFlow, NLP, Machine Learning, Deep Learning, Research, Teaching")

    doc.add_heading("Languages", level=1)
    doc.add_paragraph("English (Native), Spanish (Native), French (Conversational)")
    save_docx(doc, "05_academic.docx")


# ---------------------------------------------------------------------------
# Scenario 6: date-first PDF format (the srbhr barry pattern)
# ---------------------------------------------------------------------------
def scenario_date_first_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Liam O'Connor", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 6, "liam.oconnor@example.com | (503) 555-0157 | Portland, OR",
             new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "B", 12)
    pdf.ln(4)
    pdf.cell(0, 8, "EXPERIENCE", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    pdf.ln(1)
    pdf.cell(0, 7, "Software Engineer (Front-End), Nexus Inc, Portland, OR, USA", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(0, 5, "June 2022 - Present Developing intuitive user interfaces using React and Redux. Working closely with UX designers to implement responsive and accessible web design. Participating in agile development processes.")

    pdf.set_font("Helvetica", "B", 10)
    pdf.ln(3)
    pdf.cell(0, 7, "Front-End Developer, Brightpath, Seattle, WA, USA", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(0, 5, "August 2019 - May 2022 Building and maintaining e-commerce websites. Implementing A/B tests and improving conversion rates. Collaborating with back-end developers to integrate RESTful APIs.")

    pdf.set_font("Helvetica", "B", 10)
    pdf.ln(3)
    pdf.cell(0, 7, "Web Developer Intern, Sparkbyte, Portland, OR, USA", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(0, 5, "May 2018 - August 2018 Assisting with website maintenance and feature development.")

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "EDUCATION", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.ln(1)
    pdf.cell(0, 6, "B.Sc. Computer Science, University of Oregon, 2018", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "SKILLS", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.ln(1)
    pdf.cell(0, 6, "React, Redux, JavaScript, TypeScript, CSS, Git, Agile", new_x="LMARGIN", new_y="NEXT")

    path = os.path.join(OUT_DIR, "06_date_first.pdf")
    pdf.output(path)
    print("  wrote 06_date_first.pdf")


# ---------------------------------------------------------------------------
# Scenario 7: multi-degree paragraph education (template pattern, real dates)
# ---------------------------------------------------------------------------
def scenario_multidegree():
    doc = Document()
    doc.add_heading("Dr. Aisha Khan", level=0)
    doc.add_paragraph("aisha.khan@example.com | (212) 555-0134 | New York, NY")
    doc.add_paragraph("Data Scientist with a research background.")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Senior Data Scientist, Insightful AI | Mar 2021 - Present")
    doc.add_paragraph("Built ML models for fraud detection with Python and scikit-learn.")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("Ph.D. Computer Science")
    doc.add_paragraph("Harvard University | Sep 2014 - May 2019")
    doc.add_paragraph("Dissertation: Robust Learning under Label Noise.")
    doc.add_paragraph("M.Sc. Statistics")
    doc.add_paragraph("Cornell University | Aug 2012 - May 2014")
    doc.add_paragraph("Thesis: Bayesian Methods for High-Dimensional Data.")
    doc.add_paragraph("B.Sc. Mathematics")
    doc.add_paragraph("NYU | Aug 2008 - May 2012")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, R, scikit-learn, PyTorch, Pandas, NumPy, SQL, Statistical Analysis, Machine Learning")

    doc.add_heading("Languages", level=1)
    doc.add_paragraph("English (Native), Urdu (Native), Arabic (Conversational)")
    save_docx(doc, "07_multidegree.docx")


# ---------------------------------------------------------------------------
# Scenario 8: ORG false-positive trigger (teacher's assistant style)
# ---------------------------------------------------------------------------
def scenario_org_false_positive():
    doc = Document()
    doc.add_heading("Rachel Thompson", level=0)
    doc.add_paragraph("rachel.thompson@example.com | (512) 555-0182 | Austin, TX")
    doc.add_paragraph("Recent graduate seeking software engineering roles.")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Teacher's Assistant, University of Texas at Austin | August 2022 - Present")
    doc.add_paragraph("- Conduct class discussions and grade assignments")
    doc.add_paragraph("Software Engineer Intern, TechCorp Inc | May 2022 - Aug 2022")
    doc.add_paragraph("- Collaborated on web applications using JavaScript and React")
    doc.add_paragraph("Research Assistant, DataLab | Aug 2020 - May 2022")
    doc.add_paragraph("- Analyzed datasets with Python and Pandas")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.Sc. Computer Science, University of Texas at Austin, May 2024")

    doc.add_heading("Project Highlights", level=1)
    doc.add_paragraph("Courseable Application (Java)")
    doc.add_paragraph("Snake Game")
    doc.add_paragraph("Portfolio Website")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Java, JavaScript, React, SQL, Git, Pandas, Communication")

    doc.add_heading("Activities", level=1)
    doc.add_paragraph("Member, Women in CS Club")
    save_docx(doc, "08_org_false_positive.docx")


# ---------------------------------------------------------------------------
# Scenario 9: sparse entry-level TXT (junior pattern)
# ---------------------------------------------------------------------------
def scenario_sparse_txt():
    text = """Zoe Miller
Recent Graduate | Entry-Level Developer
zoe.miller@example.com | (720) 555-0121 | Denver, CO

EDUCATION
Bachelor of Science in Information Technology
Colorado State University | May 2025
GPA: 3.2/4.0

PROJECTS
Library Management System (Capstone)
- Built a web app with React and Node.js
- Used MongoDB for data storage

TECHNICAL SKILLS
Languages: JavaScript (basic), Python (basic), HTML, CSS
Tools: Git, VS Code, MongoDB

ACTIVITIES
- Member, Robotics Club
- Volunteer, Local Food Bank
"""
    path = os.path.join(OUT_DIR, "09_sparse_entry.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    print("  wrote 09_sparse_entry.txt")


# ---------------------------------------------------------------------------
# Scenario 10: strong senior TXT (many projects, certs, leadership)
# ---------------------------------------------------------------------------
def scenario_strong_txt():
    text = """Nathan Brooks
Senior DevOps Engineer
nathan.brooks@example.com | (404) 555-0199 | Atlanta, GA
linkedin.com/in/nathanbrooks

PROFESSIONAL SUMMARY
Senior DevOps engineer with 10 years of experience automating cloud infrastructure and CI/CD pipelines.

WORK HISTORY
Senior DevOps Engineer, CloudScale Inc | Jan 2019 - Present
- Architect Kubernetes clusters and Terraform modules
- Cut deployment time from 45 min to 6 min
DevOps Engineer, DataWorks LLC | Jun 2016 - Dec 2018
- Built CI/CD with Jenkins and GitHub Actions
- Managed AWS infrastructure with CloudFormation
Systems Administrator, Nebula Solutions | Aug 2014 - May 2016
- Administered Linux servers and Nginx

EDUCATION
B.Sc. Computer Science, Georgia Tech, 2014

TECHNICAL SKILLS
AWS, Azure, Docker, Kubernetes, Terraform, Ansible, Jenkins, GitLab CI, GitHub Actions, CI/CD, Linux, Nginx, Prometheus, Grafana, Python, Bash, Helm, ArgoCD

CERTIFICATIONS
AWS Certified Solutions Architect, 2023
Certified Kubernetes Administrator, 2022
Google Cloud Professional Cloud Architect, 2021

PROJECTS
Home Lab Kubernetes Cluster - K3s, Prometheus, Grafana
Infra-as-Code Templates - Terraform, AWS
Deployment Slack Bot - Python, GitHub Actions

LANGUAGES
English (Native)

LEADERSHIP
Chapter Lead, DevOps Meetup Atlanta
Open-source maintainer, terraform-aws-modules
"""
    path = os.path.join(OUT_DIR, "10_strong_senior.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    print("  wrote 10_strong_senior.txt")


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    scenarios = [
        scenario_clean,
        scenario_two_column_pdf,
        scenario_table_docx,
        scenario_duplicated_docx,
        scenario_academic,
        scenario_date_first_pdf,
        scenario_multidegree,
        scenario_org_false_positive,
        scenario_sparse_txt,
        scenario_strong_txt,
    ]
    for fn in scenarios:
        fn()
    print(f"\nBenchmark CVs written to {OUT_DIR}")


if __name__ == "__main__":
    main()
