"""
tests/test_parser.py
CVInsight — Week 2, Day 9

Tests for:
  - src/parser/docx_parser.py  (parse_docx)
  - src/parser/txt_parser.py   (parse_txt)
  - src/parser/parser.py       (parse_cv — unified dispatcher)

Run:
    pytest tests/test_parser.py -v
"""

import os
import sys
import tempfile
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from src.parser.docx_parser import parse_docx
from src.parser.txt_parser import parse_txt
from src.parser.parser import parse_cv, get_supported_extensions

# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------

TMP = tempfile.gettempdir()

CV_CONTENT = """John Doe
john.doe@email.com | +880-1700-000000

EDUCATION
MSc Data Science, BUET, 2021

EXPERIENCE
Data Engineer, DataCorp, 2021 - 2024
  - Built ETL pipelines handling 500K records/day
  - Reduced processing time by 40%

SKILLS
Python, SQL, Apache Spark, Airflow, Docker

CERTIFICATIONS
Google Professional Data Engineer, 2022
"""


@pytest.fixture(scope="module")
def tmp_txt_cv(tmp_path_factory):
    """Write CV_CONTENT to a .txt file."""
    p = tmp_path_factory.mktemp("data") / "cv.txt"
    p.write_text(CV_CONTENT, encoding="utf-8")
    return str(p)


@pytest.fixture(scope="module")
def tmp_docx_cv(tmp_path_factory):
    """Create a .docx CV file using python-docx."""
    pytest.importorskip("docx", reason="python-docx required")
    from docx import Document

    doc = Document()
    for line in CV_CONTENT.strip().splitlines():
        doc.add_paragraph(line)

    p = tmp_path_factory.mktemp("data") / "cv.docx"
    doc.save(str(p))
    return str(p)


@pytest.fixture(scope="module")
def tmp_docx_with_table(tmp_path_factory):
    """Create a .docx CV with a two-column skills table."""
    pytest.importorskip("docx", reason="python-docx required")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Alice Chen")
    doc.add_paragraph("alice@example.com")
    doc.add_paragraph("SKILLS")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Machine Learning"
    table.cell(1, 0).text = "TensorFlow"
    table.cell(1, 1).text = "Docker"

    p = tmp_path_factory.mktemp("data") / "cv_table.docx"
    doc.save(str(p))
    return str(p)


# ---------------------------------------------------------------------------
# TXT Parser tests
# ---------------------------------------------------------------------------

class TestParseTxt:

    def test_returns_string(self, tmp_txt_cv):
        assert isinstance(parse_txt(tmp_txt_cv), str)

    def test_non_empty(self, tmp_txt_cv):
        assert len(parse_txt(tmp_txt_cv)) > 50

    def test_name_present(self, tmp_txt_cv):
        assert "John Doe" in parse_txt(tmp_txt_cv)

    def test_email_present(self, tmp_txt_cv):
        assert "john.doe@email.com" in parse_txt(tmp_txt_cv)

    def test_skills_present(self, tmp_txt_cv):
        assert "Python" in parse_txt(tmp_txt_cv)

    def test_no_excessive_blank_lines(self, tmp_txt_cv):
        assert "\n\n\n" not in parse_txt(tmp_txt_cv)

    def test_file_not_found_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_txt(os.path.join(TMP, "nonexistent_cv.txt"))

    def test_wrong_extension_raises(self, tmp_txt_cv):
        # Rename logic: just pass a .pdf path to txt parser
        with pytest.raises(ValueError):
            parse_txt(tmp_txt_cv.replace(".txt", ".pdf"))

    def test_latin1_encoding(self, tmp_path_factory):
        """CV saved in latin-1 — only chars in latin-1 range (no em dash)."""
        p = tmp_path_factory.mktemp("enc") / "cv_latin1.txt"
        p.write_bytes("Resume - Joehn Doee\nSkills: Python\n".encode("latin-1"))
        result = parse_txt(str(p))
        assert "Python" in result

    def test_cp1252_encoding(self, tmp_path_factory):
        """CV saved in cp1252 (Windows default) with em dash and accents."""
        p = tmp_path_factory.mktemp("enc") / "cv_cp1252.txt"
        p.write_bytes("Résumé — Jöhn Döe\nSkills: Python\n".encode("cp1252"))
        result = parse_txt(str(p))
        assert "Python" in result


# ---------------------------------------------------------------------------
# DOCX Parser tests
# ---------------------------------------------------------------------------

class TestParseDocx:

    def test_returns_string(self, tmp_docx_cv):
        assert isinstance(parse_docx(tmp_docx_cv), str)

    def test_non_empty(self, tmp_docx_cv):
        assert len(parse_docx(tmp_docx_cv)) > 50

    def test_name_present(self, tmp_docx_cv):
        assert "John Doe" in parse_docx(tmp_docx_cv)

    def test_email_present(self, tmp_docx_cv):
        assert "john.doe@email.com" in parse_docx(tmp_docx_cv)

    def test_skills_present(self, tmp_docx_cv):
        assert "Python" in parse_docx(tmp_docx_cv)

    def test_no_excessive_blank_lines(self, tmp_docx_cv):
        assert "\n\n\n" not in parse_docx(tmp_docx_cv)

    def test_table_cells_extracted(self, tmp_docx_with_table):
        result = parse_docx(tmp_docx_with_table)
        assert "Python" in result
        assert "TensorFlow" in result
        assert "Machine Learning" in result

    def test_file_not_found_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_docx(os.path.join(TMP, "nonexistent_cv.docx"))

    def test_wrong_extension_raises(self, tmp_docx_cv):
        with pytest.raises(ValueError):
            parse_docx(tmp_docx_cv.replace(".docx", ".pdf"))


# ---------------------------------------------------------------------------
# Unified parse_cv dispatcher tests
# ---------------------------------------------------------------------------

class TestParseCv:

    def test_dispatches_txt(self, tmp_txt_cv):
        result = parse_cv(tmp_txt_cv)
        assert "John Doe" in result

    def test_dispatches_docx(self, tmp_docx_cv):
        result = parse_cv(tmp_docx_cv)
        assert "John Doe" in result

    def test_unsupported_extension_raises(self):
        fake = os.path.join(TMP, "cv.rtf")
        # Create a dummy file so FileNotFoundError doesn't fire first
        open(fake, "w").close()
        try:
            with pytest.raises(ValueError, match="Unsupported file type"):
                parse_cv(fake)
        finally:
            os.remove(fake)

    def test_file_not_found_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_cv(os.path.join(TMP, "ghost_cv.pdf"))

    def test_get_supported_extensions(self):
        exts = get_supported_extensions()
        assert ".pdf" in exts
        assert ".docx" in exts
        assert ".txt" in exts